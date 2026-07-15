#!/usr/bin/env python3
"""Build the professor-ready NeuralForensics technical explainer."""

from __future__ import annotations

import math
import pathlib
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = pathlib.Path(__file__).resolve().parent
OUTPUT = ROOT / "NeuralForensics_NSU_Technical_Explainer.docx"
ASSET_DIR = ROOT / ".doc_assets"

# compact_reference_guide preset, with named NSU cyber-forensics color overrides.
PAGE_WIDTH = 12240
PAGE_HEIGHT = 15840
CONTENT_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = (80, 80, 120, 120)

NAVY = "102A43"
BLUE = "167D9A"
CYAN = "00A9B7"
PALE_CYAN = "E8F7F8"
ICE = "F4F8FA"
INK = "172B3A"
MUTED = "5B7080"
GRID = "B9CBD3"
RED = "C5283D"
PALE_RED = "FCECEF"
GREEN = "158A69"
PALE_GREEN = "E8F6F1"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, *CELL_MARGINS)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_font(run, name="Aptos", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color=CYAN, size=14, space=7) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.append(begin)
    run.append(instr)
    run.append(separate)
    run.append(end)


def define_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def create(abstract_id: int, num_id: int, fmt: str, text_value: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_id))
        num.append(abs_ref)
        numbering.append(num)

    create(next_abs, next_num, "bullet", "•")
    create(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def apply_list(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def set_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("NEURALFORENSICS NSU")
    set_font(r, "Aptos", 8.5, NAVY, True)
    r = p.add_run("  /  TECHNICAL EXPLAINER")
    set_font(r, "Aptos", 8.5, MUTED, False)
    paragraph_border_bottom(p, GRID, 6, 4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("North South University Cyber Security Research Area  |  ")
    set_font(r, "Aptos", 8, MUTED)
    add_field(p, "PAGE")


def add_body(doc, text: str, bold_lead: str | None = None, italic=False) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, italic=italic)


def add_bullet(doc, text: str, bullet_id: int) -> None:
    p = doc.add_paragraph()
    apply_list(p, bullet_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))


def add_number(doc, text: str, number_id: int) -> None:
    p = doc.add_paragraph()
    apply_list(p, number_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))


def add_callout(doc, label: str, text: str, kind="info") -> None:
    fills = {"info": PALE_CYAN, "warning": PALE_RED, "success": PALE_GREEN}
    colors = {"info": BLUE, "warning": RED, "success": GREEN}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[kind])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(label.upper() + "  "), size=9, color=colors[kind], bold=True)
    set_font(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2:
            for cell in cells:
                set_cell_shading(cell, ICE)
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 18 and idx != len(values) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(value), size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0C1B2A")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        set_font(p.add_run(line), "Courier New", 8.5, "D7F7FA")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def create_architecture_diagram() -> pathlib.Path:
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "architecture.png"
    scale = 2
    image = Image.new("RGB", (1400 * scale, 620 * scale), "#F4F8FA")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 28 * scale)
        small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 19 * scale)
        bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 28 * scale)
    except OSError:
        font = small = bold = ImageFont.load_default()

    boxes = [
        (50, 175, 285, 405, "Browser UI", "HTML / CSS / JS\nCamera or MP4\nScan animation"),
        (385, 175, 650, 405, "Local HTTP Server", "Static files\nPOST /api/analyze\nMultipart routing"),
        (750, 65, 1080, 275, "Detection Model 1", "PyTorch Meso4\n12-frame sampling\nCPU inference"),
        (750, 345, 1080, 555, "Detection Model 2", "Optional remote video\nclassification service\nserver-held token"),
        (1180, 175, 1350, 405, "Verdict", "Normalized score\n0.50 threshold\nred / green UI"),
    ]
    for x1, y1, x2, y2, title, body in boxes:
        coords = tuple(v * scale for v in (x1, y1, x2, y2))
        draw.rounded_rectangle(coords, radius=18 * scale, fill="#FFFFFF", outline="#167D9A", width=3 * scale)
        draw.text(((x1 + 20) * scale, (y1 + 24) * scale), title, font=bold, fill="#102A43")
        draw.multiline_text(((x1 + 20) * scale, (y1 + 82) * scale), body, font=small, fill="#425B6B", spacing=10 * scale)

    def arrow(start, end, label=""):
        sx, sy = start[0] * scale, start[1] * scale
        ex, ey = end[0] * scale, end[1] * scale
        draw.line((sx, sy, ex, ey), fill="#00A9B7", width=5 * scale)
        angle = math.atan2(ey - sy, ex - sx)
        for delta in (2.6, -2.6):
            draw.line((ex, ey, ex + 18 * scale * math.cos(angle + delta), ey + 18 * scale * math.sin(angle + delta)), fill="#00A9B7", width=5 * scale)
        if label:
            mx, my = (sx + ex) / 2, (sy + ey) / 2
            draw.text((mx - 35 * scale, my - 30 * scale), label, font=small, fill="#167D9A")

    arrow((285, 290), (385, 290), "FormData")
    arrow((650, 260), (750, 170), "route")
    arrow((650, 320), (750, 450), "route")
    arrow((1080, 170), (1180, 260), "JSON")
    arrow((1080, 450), (1180, 330), "JSON")
    image.resize((1400, 620), Image.Resampling.LANCZOS).save(path)
    return path


def page_break(doc) -> None:
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("NORTH SOUTH UNIVERSITY"), "Aptos", 11, CYAN, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_font(p.add_run("NeuralForensics NSU"), "Aptos Display", 31, NAVY, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run("Deepfake Detection Web Application"), "Aptos Display", 18, BLUE, False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(54)
    set_font(p.add_run("Complete Technical Explainer, Implementation Record, and Faculty Presentation Guide"), "Aptos", 12, MUTED, False, True)

    table = doc.add_table(rows=3, cols=2)
    values = [
        ("Research area", "Cyber Security"),
        ("Application", "NeuralForensics NSU // Cyber Lab"),
        ("Prepared", date.today().strftime("%d %B %Y")),
    ]
    for row, (label, value) in zip(table.rows, values):
        set_cell_shading(row.cells[0], PALE_CYAN)
        set_font(row.cells[0].paragraphs[0].add_run(label.upper()), size=9, color=BLUE, bold=True)
        set_font(row.cells[1].paragraphs[0].add_run(value), size=10.5, color=INK)
    set_table_geometry(table, [2700, 6660])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Professor-ready edition"), size=10, color=GREEN, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Built with HTML5, CSS3, vanilla JavaScript, Python, PyTorch, and FFmpeg"), size=9.5, color=MUTED)


def build() -> pathlib.Path:
    doc = Document()
    configure_styles(doc)
    set_running_furniture(doc.sections[0])
    bullet_id, number_id = define_numbering(doc)
    architecture = create_architecture_diagram()

    core = doc.core_properties
    core.title = "NeuralForensics NSU Technical Explainer"
    core.subject = "Deepfake detection application architecture, implementation, testing, and presentation guide"
    core.author = "North South University Cyber Security Research Area"
    core.keywords = "deepfake, Meso4, MesoNet, video forensics, PyTorch, NSU"

    add_cover(doc)
    page_break(doc)

    doc.add_heading("How to Use This Document", level=1)
    add_body(doc, "This document explains the complete project as it currently exists in the codebase. It is designed both as a technical reference and as a speaking guide for a faculty demonstration. It separates implemented facts, optional integrations, measured test results, and known limitations so that the project can be presented accurately.")
    add_callout(doc, "Terminology correction", "The implemented local network is Meso4, one of the compact MesoNet architectures. If the project has been called the “MISA model” during discussion, that name should be corrected to Meso4. There is no MISA model class or MISA weight file in this repository.", "warning")
    doc.add_heading("Reading Map", level=2)
    for text in [
        "Sections 1–3 explain the problem, user experience, and full architecture.",
        "Sections 4–7 explain the browser, backend, Meso4 network, preprocessing, inference, and dual-model routing.",
        "Sections 8–10 explain the API path, security, testing, failures, and limitations.",
        "Sections 11–14 provide deployment guidance, a demonstration runbook, a professor-facing speaking script, viva questions, and a glossary.",
    ]:
        add_bullet(doc, text, bullet_id)
    doc.add_heading("Executive Summary", level=2)
    add_body(doc, "NeuralForensics NSU // Cyber Lab is a zero-build browser interface connected to a Python inference server. A user can activate a webcam only after explicit permission, record a three-second clip, upload an MP4, or load one of three bundled known-synthetic test videos. The server routes the media to one of two selectable detection paths and returns a normalized probability. The browser converts that probability into a clear red or green verdict at a threshold of 0.50.")
    add_body(doc, "Detection Model 1 is a compact Meso4 convolutional neural network executed locally with PyTorch on the CPU. It samples up to 12 frames, center-crops and normalizes them, classifies each frame, and averages the frame scores. Detection Model 2 is an optional remote video-classification path whose credential remains on the server. Both paths return the same response shape, allowing the user interface to remain implementation-neutral.")
    add_callout(doc, "Central scientific point", "The percentage displayed by the interface is synthetic likelihood or model confidence for this input. It is not the empirical accuracy of the system. Accuracy requires a labeled evaluation dataset and metrics such as ROC-AUC, sensitivity, specificity, precision, recall, and false-positive rate.")

    page_break(doc)
    doc.add_heading("1. Project Purpose and Scope", level=1)
    doc.add_heading("1.1 Problem being addressed", level=2)
    add_body(doc, "Deepfakes are manipulated or AI-generated media in which facial identity, expression, or visual content can be synthesized. The project demonstrates how a browser-based forensic workflow can collect video evidence, send it to a detection engine, and communicate a risk verdict in an understandable way. Its purpose is educational demonstration and research prototyping, not legal certification or autonomous enforcement.")
    doc.add_heading("1.2 Functional objectives", level=2)
    for text in [
        "Provide two obvious inputs: an explicitly authorized live camera and a local MP4 upload.",
        "Keep the camera off until the user deliberately enables it.",
        "Make analysis visually legible through an active scanner, terminal messages, and an in-viewport result.",
        "Run a bundled local detector without requiring internet connectivity.",
        "Offer a second selectable detector through the same interface without exposing provider details in the user-facing UI.",
        "Include known-synthetic videos inside the project so professors and users can verify the end-to-end workflow.",
        "Use standard HTML5, CSS3, and vanilla JavaScript so the frontend requires no compilation or package build step.",
    ]:
        add_bullet(doc, text, bullet_id)
    doc.add_heading("1.3 What the system does not claim", level=2)
    for text in [
        "It does not prove that a video is genuine; it estimates whether learned visual patterns resemble the model’s fake class.",
        "It does not identify who created a manipulation or which exact generation tool was used.",
        "It does not provide calibrated production accuracy across every camera, codec, ethnicity, lighting condition, or manipulation method.",
        "It does not perform audio deepfake detection.",
        "It does not currently store a case history, user account, audit database, or cryptographic chain of custody.",
    ]:
        add_bullet(doc, text, bullet_id)

    doc.add_heading("2. User Experience and Interface", level=1)
    doc.add_heading("2.1 Visual system", level=2)
    add_body(doc, "The frontend uses a dark charcoal-to-midnight gradient, translucent glass panels, cyan forensic accents, red alert accents, a monospace display font, an animated grid, moving aurora glows, concentric biometric rings, and a stylized face mesh. These elements create a cyber-forensics command-center identity while preserving a conventional input → scan → result workflow.")
    add_table(doc, ["UI element", "Implementation", "Purpose"], [
        ["Glass panels", "RGBA panel, 16 px backdrop blur, cyan border, layered shadow", "Separates functional areas while retaining depth"],
        ["Scanner line", "CSS keyframe animation; visible only during analysis", "Shows that processing is active"],
        ["Terminal overlay", "JavaScript typewriter loop with three forensic messages", "Makes asynchronous wait time understandable"],
        ["Targeting HUD", "CSS lines, nodes, corners, and telemetry", "Guides face placement and strengthens visual identity"],
        ["Verdict state", "Red alert or cyan/green authentic state", "Communicates the result without requiring technical interpretation"],
    ], [1900, 3600, 3860])
    doc.add_heading("2.2 Input hierarchy", level=2)
    add_body(doc, "The right-side control panel presents numbered choices. Detection Model 1 or Detection Model 2 is chosen first. The user then enables the camera or selects an MP4. These sources are mutually exclusive: loading a file stops the camera, and enabling the camera clears a selected file. The large action button beneath the viewport changes its label according to state and remains disabled until an input is ready.")
    doc.add_heading("2.3 Verdict location", level=2)
    add_body(doc, "The result panel is positioned inside the media viewport. During scanning it fades away so the video, laser, and terminal overlay remain visible. After processing, it returns with the input name, synthetic score, selected model, headline, and explanation. This placement was adopted because an earlier separate status area made input and output difficult to associate.")

    page_break(doc)
    doc.add_heading("3. End-to-End Architecture", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(architecture), width=Inches(6.45))
    p = doc.add_paragraph("Figure 1. Current request and inference architecture")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.runs[0], size=9, color=MUTED, italic=True)
    doc.add_heading("3.1 Component inventory", level=2)
    add_table(doc, ["Component", "Technology", "Responsibility"], [
        ["index.html", "HTML5", "Semantic dashboard structure, video elements, controls, model selector, sample library"],
        ["style.css", "CSS3", "Responsive layout, glassmorphism, animations, scanner, HUD, state colors"],
        ["app.js", "Vanilla JavaScript", "State, camera permission, recording, drag/drop, requests, parsing, verdict rendering"],
        ["detector_server.py", "Python HTTP server", "Static hosting, multipart endpoint, routing, decoding, model inference, API proxy"],
        ["meso4_best.pth", "PyTorch weights", "Learned parameters for the local two-class Meso4 detector"],
        ["NSU_Demo_Videos", "MP4 assets", "Three known-synthetic verification inputs exposed in the interface"],
        ["fetch_test_media.py", "Python utility", "Generates illustrative local MP4 clips; not the source of the three actual deepfake samples"],
    ], [1950, 1850, 5560], 9.1)
    doc.add_heading("3.2 Normalized detector contract", level=2)
    add_body(doc, "The frontend does not need provider-specific response logic. Both model routes are converted to a common project response with status, a 0–1 synthetic probability, and metadata. This abstraction is what makes two detectors interchangeable from the browser’s perspective.")
    add_code(doc, '{\n  "status": "success",\n  "type": {"ai_generated": 0.61, "deepfake": 0.61},\n  "meta": {\n    "detector": "Detection Model 1 or 2",\n    "source_type": "video or live_video",\n    "frame_count": 12,\n    "max_frame_score": 0.66,\n    "mean_frame_score": 0.61\n  }\n}')

    doc.add_heading("4. Browser-Side Implementation", level=1)
    doc.add_heading("4.1 Application state", level=2)
    add_body(doc, "app.js keeps a small in-memory state object containing the active MediaStream, selected model, selected file, object URL, scan lock, and terminal timers. This avoids a framework while still enforcing predictable transitions. The scan lock prevents duplicate requests, and cleanupMedia stops camera tracks and revokes object URLs when the page closes.")
    doc.add_heading("4.2 Camera authorization and privacy", level=2)
    add_body(doc, "The browser calls navigator.mediaDevices.getUserMedia only after the Enable Camera button is clicked. It requests a user-facing video track with an ideal resolution of 1280 × 720 and explicitly disables audio. If permission is denied or no camera exists, the error is converted into a readable status and the MP4 workflow remains available.")
    add_callout(doc, "Secure-context requirement", "Camera APIs require a secure context. localhost is treated as trustworthy for development, while a deployed site must use HTTPS. Opening index.html through file:// is not a reliable camera execution path.", "warning")
    doc.add_heading("4.3 Live clip acquisition", level=2)
    for text in [
        "The camera preview is a live MediaStream displayed in the webcam <video> element.",
        "When analysis starts, MediaRecorder checks VP8 WebM support and records the same stream for 3,000 ms.",
        "Data is collected in 250 ms chunks and combined into one WebM Blob.",
        "The Blob is appended to FormData as live-camera-scan.webm with source_type=live_video and the selected engine.",
        "The browser POSTs the multipart body to /api/analyze and waits for normalized JSON.",
    ]:
        add_number(doc, text, number_id)
    doc.add_heading("4.4 MP4 acquisition", level=2)
    add_body(doc, "An MP4 can be chosen through the hidden file input or dropped onto the labeled drop zone. JavaScript checks the MIME type or .mp4 extension, creates a temporary object URL for playback, updates metadata, and enables the analysis command. Automatic analysis is deliberately disabled, so selecting a file does not silently transmit or process it.")
    doc.add_heading("4.5 Verification sample library", level=2)
    add_body(doc, "Each sample button points to an MP4 in NSU_Demo_Videos. The browser fetches the same-origin asset, wraps its Blob in a File object, and sends it through the ordinary upload workflow. This is important: the samples do not use a privileged or hard-coded verdict path. They exercise the same decoder, preprocessing, model, response, threshold, and interface logic as a user upload.")
    doc.add_heading("4.6 Scan animation lifecycle", level=2)
    add_body(doc, "startScanEffects adds the is-scanning class and begins a character-by-character terminal sequence. CSS responds by activating the laser and hiding the previous result. stopScanEffects always runs in the finally block, whether the request succeeds or fails, preventing a permanently spinning interface after an exception.")

    doc.add_heading("5. Python Server and Infrastructure", level=1)
    doc.add_heading("5.1 Integrated static and inference server", level=2)
    add_body(doc, "detector_server.py uses ThreadingHTTPServer and SimpleHTTPRequestHandler. The same process serves index.html, JavaScript, CSS, model assets, and sample videos, and accepts POST requests at /api/analyze. This avoids cross-origin configuration during local use and supports multiple independent HTTP requests through threads.")
    doc.add_heading("5.2 Request handling", level=2)
    for text in [
        "Reject any POST path other than /api/analyze.",
        "Require multipart/form-data and parse media, engine, and source_type fields.",
        "Route legacy image batches, uploaded/live videos, and Model 2 requests separately.",
        "Read the upload into memory; Model 1 then writes supported videos to a temporary file because FFmpeg expects a path.",
        "Decode and classify the temporary video, delete it in a finally block, and serialize JSON.",
        "Catch errors and return a failure object with HTTP 500 so the browser can show a scan failure rather than inventing a result.",
    ]:
        add_number(doc, text, number_id)
    doc.add_heading("5.3 Dependencies", level=2)
    add_table(doc, ["Dependency", "Role", "Operational note"], [
        ["PyTorch", "Defines Meso4, loads weights, performs tensor inference", "CPU-only use is sufficient for this compact model"],
        ["NumPy", "Frame indexing, normalization, aggregation, image arrays", "Used throughout preprocessing"],
        ["Pillow", "RGB conversion, center crop output, 256 × 256 resize", "Bilinear interpolation"],
        ["imageio + FFmpeg", "Reads MP4/WebM/MOV and extracts frames", "FFmpeg binary is provided in the project vendor path"],
        ["Python stdlib HTTP", "Static server, request handler, remote URL request", "No Flask or FastAPI runtime required"],
    ], [1700, 3760, 3900])
    doc.add_heading("5.4 Current operational topology", level=2)
    add_body(doc, "The server binds to 0.0.0.0 on port 8080. Binding to 0.0.0.0 permits network interfaces to accept connections, while the printed development address is http://localhost:8080. There is no database, queue, persistent upload storage, user identity layer, or container manifest in the current implementation.")

    doc.add_heading("6. Detection Model 1: Meso4", level=1)
    doc.add_heading("6.1 What Meso4 is", level=2)
    add_body(doc, "Meso4 is a compact convolutional neural network from the MesoNet family introduced by Afchar, Nozick, Yamagishi, and Echizen. “Meso” refers to mesoscopic visual properties: patterns between very low-level pixel noise and high-level semantic identity. The design aims to notice manipulation artifacts in facial imagery while remaining much smaller than heavyweight image-classification backbones.")
    add_body(doc, "The original work studied Deepfake and Face2Face manipulations. The repository uses a Meso4 implementation and weight checkpoint obtained through the DeepfakeBench ecosystem. DeepfakeBench provides standardized detector implementations, dataset management, and evaluation protocols; it also demonstrates that cross-dataset performance can be substantially harder than results on a detector’s familiar training distribution.")
    doc.add_heading("6.2 Exact implemented network", level=2)
    add_table(doc, ["Stage", "Operation", "Output concept", "Why it exists"], [
        ["Input", "3-channel RGB, 256 × 256", "Color face/video crop", "Fixed tensor geometry"],
        ["Block 1", "Conv 3×3, 3→8; ReLU; BN; max-pool 2×2", "8 feature maps", "Early edges and texture"],
        ["Block 2", "Conv 5×5, 8→8; ReLU; BN; max-pool 2×2", "8 feature maps", "Larger local artifact patterns"],
        ["Block 3", "Conv 5×5, 8→16; ReLU; BN; max-pool 2×2", "16 feature maps", "Richer mesoscopic patterns"],
        ["Block 4", "Conv 5×5, 16→16; ReLU; BN; max-pool 4×4", "16 × 8 × 8", "Strong spatial compression"],
        ["Classifier", "Flatten; dropout 0.5; FC 1024→16; LeakyReLU", "16-dimensional representation", "Compact decision layer"],
        ["Output", "Dropout 0.5; FC 16→2; softmax", "P(real), P(fake)", "Binary class probabilities"],
    ], [1200, 3400, 2050, 2710], 8.8)
    add_callout(doc, "Class mapping", "The loaded DeepfakeBench checkpoint uses class index 0 for real and class index 1 for fake. The implementation applies softmax to the two logits and reports softmax(logits)[1] as synthetic likelihood.")
    doc.add_heading("6.3 Weight loading", level=2)
    add_body(doc, "The .models/meso4_best.pth checkpoint is loaded onto the CPU. Checkpoint keys beginning with backbone. are retained, that prefix is removed, and the resulting state dictionary is loaded with strict=True. Strict loading is valuable because startup fails immediately if tensor names or shapes do not match the implemented architecture. The model is switched to evaluation mode, which disables dropout randomness and uses stored batch-normalization statistics.")
    doc.add_heading("6.4 Why the model is lightweight", level=2)
    add_body(doc, "The checkpoint is approximately 115 KB and the network has only four convolutional stages plus two small fully connected layers. That makes CPU inference practical for a local demonstration. The tradeoff is lower representational capacity and weaker robustness to manipulations, cameras, codecs, and datasets unlike the training distribution.")

    doc.add_heading("7. Preprocessing, Inference, and Verdict Logic", level=1)
    doc.add_heading("7.1 Video frame sampling", level=2)
    add_body(doc, "imageio opens the uploaded or recorded video through FFmpeg. The server reads the declared frame count; when that value is missing or infinite, it estimates frame count as duration × frames per second. It creates up to 12 evenly spaced indices with NumPy linspace. This covers the whole timeline more effectively than reading only the first 12 frames while keeping latency bounded.")
    doc.add_heading("7.2 Per-frame preprocessing", level=2)
    for text in [
        "Convert the decoded frame to an RGB array.",
        "Take a centered square whose side is the smaller of frame width and height.",
        "Resize that square to 256 × 256 pixels with bilinear interpolation.",
        "Convert pixel values from 0–255 integers to 0–1 floating point.",
        "Normalize each channel with (value − 0.5) / 0.5, producing an approximate −1 to +1 range.",
        "Transpose height × width × channel into channel × height × width and add a batch dimension.",
    ]:
        add_number(doc, text, number_id)
    add_callout(doc, "Important mismatch", "DeepfakeBench preprocessing commonly uses detected, aligned face crops. This application currently uses a center square, not a face detector and landmark aligner. If the face is small or off-center, background pixels dominate and reliability can fall. Adding RetinaFace, MediaPipe Face Detection, or another alignment stage is the highest-value model-pipeline improvement.", "warning")
    doc.add_heading("7.3 Frame classification and video aggregation", level=2)
    add_body(doc, "Each prepared frame is classified independently. Model 1 stores all readable fake-class probabilities and computes their arithmetic mean as the video score. It also reports the maximum and mean frame scores as metadata. Failed individual frame reads are skipped; the request fails only when no readable frame remains.")
    add_code(doc, "frame_scores = [Meso4(frame_i).softmax()[fake] for frame_i in sampled_frames]\nvideo_score = mean(frame_scores)\nverdict = 'deepfake suspected' if video_score >= 0.50 else 'authentic signal'")
    doc.add_heading("7.4 Threshold", level=2)
    add_body(doc, "The project threshold is 0.50. At or above 0.50, the browser enters the alert state; below 0.50, it enters the authentic state. This threshold is intuitive for a binary softmax but has not been calibrated against a held-out validation set for this exact application. A production threshold must be selected from ROC or precision–recall analysis according to the cost of false positives and false negatives.")
    doc.add_heading("7.5 Live-camera correction", level=2)
    add_body(doc, "Local tests showed a domain shift: genuine webcam WebM could receive a raw Meso4 score near 0.68. For Model 1 live_video results only, the frontend applies a fixed log-odds offset before thresholding:")
    add_code(doc, "p = clip(raw_score, 0.0001, 0.9999)\nlogit = ln(p / (1 - p))\ncorrected_score = sigmoid(logit - 1.2)")
    add_body(doc, "For example, a raw score of 0.68 becomes approximately 0.39. This is a pragmatic presentation calibration, not learned calibration and not a scientifically validated correction. Uploaded MP4 scores and Model 2 scores are not changed. A defensible replacement would fit Platt scaling or isotonic regression on labeled webcam clips captured across devices and environments.")
    add_callout(doc, "Presentation language", "Say: “The live path includes an experimental camera-domain calibration because WebM compression and unaligned webcam framing differ from the model’s training distribution.” Do not say: “We changed the score so the camera would look real.”", "warning")
    doc.add_heading("7.6 Unused compatibility code", level=2)
    add_body(doc, "The repository retains a still-frame canvas capture helper and a legacy multi-image live route with brightness, contrast, detail, and temporal-change calculations. The current user flow does not call those functions; it records a three-second WebM instead. They are compatibility remnants, not part of the active verdict pipeline.")

    doc.add_heading("8. Detection Model 2 and API Integration", level=1)
    doc.add_heading("8.1 Purpose of the second route", level=2)
    add_body(doc, "Detection Model 2 provides an alternative detector without changing the interface. The browser sends engine=model2 to the local backend. The backend, not the browser, authenticates to the remote service and converts its provider-specific result to the common response shape.")
    doc.add_heading("8.2 Credential handling", level=2)
    add_body(doc, "The service token is read from the HIVE_API_KEY environment variable when a Model 2 request arrives. It is placed in the outbound Authorization header and is never embedded in index.html or app.js. This prevents visitors from viewing the credential in browser source or network request bodies sent to the local server.")
    add_code(doc, 'export HIVE_API_KEY="your-token"\npython3 detector_server.py')
    doc.add_heading("8.3 Outbound request", level=2)
    add_body(doc, "The backend constructs a multipart body with a random boundary, sets video/webm or video/mp4 according to the filename, and POSTs the media field to the configured synchronous task endpoint. It uses Python’s urllib with a 120-second timeout. HTTP rejection details are shortened and returned as an honest application error.")
    doc.add_heading("8.4 Response interpretation", level=2)
    add_body(doc, "The parser traverses status → response → output → classes. It first collects values whose class is deepfake and falls back to ai_generated only if no deepfake values exist. The current project uses the maximum returned score as the verdict probability and separately calculates their mean. This conservative rule reacts to the most suspicious detected face or frame.")
    add_callout(doc, "Provider threshold context", "Hive’s documentation recommends a deepfake-video rule based on 0.50 on consecutive frames or a percentage of frames. The current implementation uses max(score) ≥ 0.50, which is simpler and more sensitive but can increase false positives. The server should preserve timestamps and implement the consecutive-frame rule for a more faithful production integration.", "warning")
    doc.add_heading("8.5 Configuration status", level=2)
    add_body(doc, "No HIVE_API_KEY is present in the current local environment, so Model 2 has not been end-to-end validated in this workspace. Selecting it correctly returns “Detection Model 2 is not configured on this server” rather than a random or simulated answer. Trial availability, pricing, rate limits, retention, and terms are controlled by the external provider and can change.")
    doc.add_heading("8.6 Earlier Sightengine attempt", level=2)
    add_body(doc, "An earlier version targeted Sightengine’s video detection endpoint. Valid credentials were supplied, but the account returned a usage-limit response because video detection was unavailable under that plan. An image-frame fallback then produced very low values for the known deepfake videos, demonstrating that substituting isolated image moderation for a proper video detector was not reliable. The project therefore adopted the bundled Meso4 path for offline operation and retained a separate optional video-service route.")

    page_break(doc)
    doc.add_heading("9. Security, Privacy, and Data Handling", level=1)
    add_table(doc, ["Area", "Current control", "Remaining risk / improvement"], [
        ["Camera consent", "getUserMedia runs only after a button click; browser permission required", "Explain camera indicator and stop behavior during demonstrations"],
        ["Audio", "audio:false; no microphone requested", "None for current scope"],
        ["Local media", "Model 1 processes through the local Python process", "Upload bytes still enter server memory and a temporary file"],
        ["Temporary storage", "Temporary video is deleted in finally", "Crash or forced termination may leave OS temp artifacts"],
        ["External transmission", "Only Model 2 sends media to a third party", "Requires explicit disclosure, consent, retention review, and privacy policy"],
        ["API secret", "Stored in server environment, not frontend", "Production secret manager and rotation are preferable"],
        ["Upload validation", "Extension/MIME checks and decoder errors", "No explicit size limit, malware scan, rate limit, or content-length cap"],
        ["Transport", "localhost during development", "Deployment requires HTTPS; external API already uses HTTPS"],
        ["Authentication", "None", "Public deployment can be abused and consume compute or API quota"],
    ], [1700, 3650, 4010], 8.7)
    doc.add_heading("9.1 Threat model summary", level=2)
    add_body(doc, "The main operational threats are oversized uploads, malformed video decoder inputs, unauthorized consumption of Model 2 quota, API-key leakage through poor server configuration, denial of service, and over-trust in a probabilistic verdict. A production system should add reverse-proxy limits, authentication, rate limiting, strict MIME inspection, isolated decoding, security headers, structured logging, consent text, and a retention policy.")
    doc.add_heading("9.2 Ethical use", level=2)
    add_body(doc, "A deepfake score can affect reputation and should never be the sole basis for accusation or disciplinary action. Results should be treated as triage evidence and combined with provenance metadata, source verification, reverse search, frame-level review, audio analysis, and human judgment. Bias and domain-shift testing should include diverse skin tones, ages, lighting conditions, camera hardware, and compression levels.")

    doc.add_heading("10. Testing, Evidence, and Self-Correction", level=1)
    doc.add_heading("10.1 Verified sample results", level=2)
    add_body(doc, "The three actual videos in NSU_Demo_Videos were analyzed with Detection Model 1 using 12 sampled frames. All three were known synthetic and exceeded the 0.50 project threshold.")
    add_table(doc, ["Known-synthetic sample", "Mean score", "Maximum frame", "Project verdict"], [
        ["Facial interaction / hugging", "67.86%", "72.57%", "Deepfake suspected"],
        ["Meeting sequence", "55.63%", "58.33%", "Deepfake suspected"],
        ["Motion / hallway sequence", "60.97%", "63.27%", "Deepfake suspected"],
    ], [3730, 1700, 1800, 2130])
    add_callout(doc, "Interpretation", "This is a functional smoke test on three positive examples, not an accuracy study. There are no matched genuine controls here, and three samples cannot estimate sensitivity, specificity, or generalization.")
    doc.add_heading("10.2 Functional verification", level=2)
    for text in [
        "JavaScript syntax passed node --check; Python source compiled successfully.",
        "The local server returned a successful Model 1 response for Sample 03 at approximately 60.97%.",
        "A known sample transcoded to WebM was decoded through the live-video server path and produced 12 frame scores.",
        "Browser inspection confirmed both model selectors, all three sample controls, enabled MP4 analysis, and no console errors.",
        "At a 1280-pixel viewport, the dashboard had no horizontal overflow and maintained a two-column layout.",
        "Model 2 without a token returned a configuration failure instead of a fabricated result.",
    ]:
        add_bullet(doc, text, bullet_id)
    doc.add_heading("10.3 Problems encountered and corrections", level=2)
    add_table(doc, ["Observed problem", "Root cause", "Correction"], [
        ["Port 8080 initially unreachable", "No server process listening", "Use detector_server.py, which serves both frontend and API on 8080"],
        ["Buttons and workflow felt unclear", "Inputs, action, and result were visually separated", "Numbered source cards, dynamic primary command, in-viewport verdict"],
        ["Cartoon test videos reduced trust", "Utility generated illustrative synthetic graphics", "Replaced website library with three actual known-deepfake MP4s"],
        ["Sightengine request rejected", "Plan did not include video analysis", "Local Meso4 became the dependable offline path"],
        ["Still-image service scores were very low", "Image fallback did not represent temporal video detection", "Analyze videos through Meso4 frame sampling"],
        ["Genuine webcam scored near 68% fake", "Domain shift, WebM compression, unaligned crop", "Experimental live-only logit correction; document limitation"],
    ], [2450, 3100, 3810], 8.7)
    doc.add_heading("10.4 Required formal evaluation", level=2)
    add_body(doc, "A professor may correctly ask for a confusion matrix. The next evaluation should use a labeled set containing both real and fake videos that are not used for training, stratified by codec, resolution, manipulation family, demographics, and capture source. Record the raw score for each video, then calculate ROC-AUC, PR-AUC, equal-error rate, false-positive rate, false-negative rate, sensitivity, specificity, precision, recall, F1, and calibration error. Choose the deployment threshold only after this analysis.")

    doc.add_heading("11. Deployment and Operations", level=1)
    doc.add_heading("11.1 Local runbook", level=2)
    add_code(doc, "cd \"/Users/nusaibaalmollick/deepfake 2.0\"\npython3 detector_server.py\n# Open http://localhost:8080")
    add_body(doc, "Starting python -m http.server is insufficient because it serves static files but does not implement POST /api/analyze or load Meso4. The integrated detector_server.py process is required.")
    doc.add_heading("11.2 Cloud constraints", level=2)
    add_body(doc, "The frontend is static, but the complete application is not purely static. Model 1 requires Python, PyTorch, Pillow, NumPy, imageio, FFmpeg, the checkpoint, temporary disk, and enough RAM to decode video. Model 2 still requires a backend to protect its credential. Therefore, a static host alone can display the interface but cannot execute the current detector.")
    add_table(doc, ["Deployment pattern", "Suitability", "Tradeoff"], [
        ["Local laptop", "Best current demonstration path", "No public sharing; dependable offline Model 1"],
        ["Container host", "Best full-stack cloud fit", "Needs sufficient RAM/CPU and may not be free"],
        ["Static frontend + separate backend", "Good public architecture", "CORS, two deployments, backend cost and scaling"],
        ["Browser-side ONNX/WebAssembly", "Potential static-only future", "Requires model conversion, FFmpeg/frame extraction redesign, performance testing"],
        ["Serverless function", "Poor fit for current large uploads and PyTorch", "Body limits, cold starts, package size, temporary execution limits"],
    ], [2500, 3000, 3860])
    doc.add_heading("11.3 Production hardening checklist", level=2)
    for text in [
        "Read host and port from environment variables instead of hardcoding 8080.",
        "Replace cgi.FieldStorage, deprecated in modern Python, with a maintained web framework or multipart parser.",
        "Pin dependency versions and add a reproducible requirements file or container image.",
        "Add upload-size limits, authentication, rate limits, timeouts, structured logs, health checks, and isolated workers.",
        "Add a true face detector/alignment stage and batch inference for performance.",
        "Store no media by default; define explicit retention and consent rules for external analysis.",
        "Use a secret manager for Model 2 and disable it when no approved provider agreement exists.",
        "Build an automated test suite for routing, malformed uploads, score extraction, threshold boundaries, and browser workflows.",
    ]:
        add_bullet(doc, text, bullet_id)

    doc.add_heading("12. Demonstration Script", level=1)
    doc.add_heading("12.1 Pre-demonstration checklist", level=2)
    for text in [
        "Start detector_server.py and verify http://localhost:8080 before entering presentation mode.",
        "Keep Detection Model 1 selected for dependable offline inference.",
        "Confirm the three verification videos load and Sample 03 completes successfully.",
        "Test camera permission in the exact browser and presentation account; then disable the camera again.",
        "Close unrelated tabs and avoid claiming the score is accuracy.",
    ]:
        add_bullet(doc, text, bullet_id)
    doc.add_heading("12.2 Suggested five-minute explanation", level=2)
    add_number(doc, "Purpose: “NeuralForensics NSU is a research prototype that performs video-forensics triage through a local compact neural detector and an optional second detector.”", number_id)
    add_number(doc, "Architecture: “The browser captures a user-approved three-second WebM or accepts an MP4, then posts it to one local endpoint. The backend selects the requested model and normalizes the result.”", number_id)
    add_number(doc, "Model: “Detection Model 1 is Meso4. It learns mesoscopic facial artifact patterns using four convolutional stages and returns fake-class probability for each sampled frame.”", number_id)
    add_number(doc, "Video logic: “We sample twelve evenly distributed frames, normalize each to the model’s expected range, infer every frame, and average the fake probabilities.”", number_id)
    add_number(doc, "Verification: “These three buttons load known-synthetic videos through the exact same user upload pipeline; their results are not hard-coded.”", number_id)
    add_number(doc, "Limitation: “The score is confidence, not accuracy. Our webcam path has an experimental domain calibration because the model was not trained for every camera and codec.”", number_id)
    add_number(doc, "Future work: “The next research step is face alignment, a balanced held-out evaluation set, formal threshold calibration, and an ensemble with stronger cross-dataset detectors.”", number_id)
    doc.add_heading("12.3 Live demonstration order", level=2)
    for text in [
        "Select Sample 03 and show that VIDEO READY appears.",
        "Click Run MP4 Deepfake Analysis; point out the active scanner and terminal overlay.",
        "Explain the red verdict, 12-frame sampling, and synthetic likelihood.",
        "Enable the camera only after explaining consent; center the face and record the three-second clip.",
        "Show that the camera can be disabled immediately and that no audio was requested.",
        "If Model 2 is not configured, do not select it during the live demonstration; explain it architecturally instead.",
    ]:
        add_number(doc, text, number_id)

    page_break(doc)
    doc.add_heading("13. Likely Professor Questions and Defensible Answers", level=1)
    qa = [
        ("Why Meso4?", "It is compact, fast on CPU, easy to bundle, and appropriate for demonstrating an end-to-end forensic pipeline. It is not claimed to be the strongest modern detector."),
        ("What is the difference between MISA and Meso4?", "The code implements Meso4. “MISA” was an informal naming error and should not be used in the technical presentation."),
        ("Is 67% the model’s accuracy?", "No. It is the fake-class probability for that input after aggregation. Accuracy is measured across a labeled dataset."),
        ("Why average 12 frames?", "Even spacing covers the timeline with bounded CPU cost. Averaging reduces dependence on one anomalous frame, although better temporal models could improve this."),
        ("Does the network analyze motion?", "Meso4 classifies frames independently. The application samples across time, but Model 1 has no recurrent, 3D-convolution, or transformer temporal module."),
        ("Why can a real webcam trigger a high raw score?", "The model expects training-like face crops. Webcam framing, lighting, WebM compression, sensor noise, and background produce domain shift. The current correction is experimental."),
        ("Are the bundled sample verdicts hard-coded?", "No. The buttons fetch MP4 files and pass them through the normal upload and inference route."),
        ("What happens without internet?", "Model 1 remains functional because inference and assets are local. Model 2 requires network access and a configured credential."),
        ("Why hide model names in the UI?", "The product interface presents stable Detection Model 1/2 labels. The technical documentation still discloses implementation details for transparency."),
        ("Can this be used as court evidence?", "Not by itself. It is a research triage prototype without validated forensic error rates, chain of custody, provenance controls, or accreditation."),
        ("What would improve it most?", "Face detection/alignment, calibrated evaluation on diverse real/fake data, stronger cross-domain detectors, temporal modeling, and explainability such as artifact heatmaps."),
        ("How is the API key protected?", "It exists only as a backend environment variable and is added to the server-to-provider Authorization header, never sent to the browser."),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        set_keep_with_next(p)
        set_font(p.add_run("Q. " + question), color=BLUE, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        set_font(p.add_run("A. " + answer), color=INK)

    doc.add_heading("14. Recommended Research Roadmap", level=1)
    add_table(doc, ["Priority", "Improvement", "Research value"], [
        ["P0", "Build balanced labeled evaluation set and report full metrics", "Makes performance claims scientifically defensible"],
        ["P0", "Add face detection and landmark alignment", "Reduces background and framing domain shift"],
        ["P1", "Calibrate separate thresholds for MP4 and webcam domains", "Controls false-positive/false-negative tradeoffs"],
        ["P1", "Implement frame timestamps and consecutive-frame decision rules", "Uses temporal evidence more correctly"],
        ["P1", "Benchmark newer DeepfakeBench detectors and ensemble outputs", "Improves cross-dataset robustness"],
        ["P2", "Add Grad-CAM or artifact-region visualization", "Improves interpretability for reviewers"],
        ["P2", "Add provenance/C2PA and metadata checks", "Combines learned detection with origin evidence"],
        ["P2", "Add audio deepfake and audio-video synchronization analysis", "Extends beyond visual-only manipulation"],
    ], [1050, 4200, 4110], 8.8)

    doc.add_heading("15. Glossary", level=1)
    add_table(doc, ["Term", "Meaning in this project"], [
        ["Deepfake", "AI-assisted synthetic or manipulated facial media."],
        ["MesoNet", "A family of compact CNNs designed to learn mesoscopic forgery artifacts."],
        ["Meso4", "The four-convolution-stage local detector implemented as Detection Model 1."],
        ["CNN", "Convolutional neural network; learns spatial filters from image tensors."],
        ["Logit", "Unnormalized model output before softmax; log-odds are also called logits in calibration."],
        ["Softmax", "Transforms two output logits into probabilities that sum to one."],
        ["Domain shift", "A difference between training data and real use, such as webcam WebM versus aligned dataset faces."],
        ["Calibration", "Adjusting scores so predicted probabilities better match observed frequencies."],
        ["Threshold", "Decision boundary that maps a continuous score to an alert or authentic state."],
        ["False positive", "A real video incorrectly flagged as fake."],
        ["False negative", "A fake video incorrectly classified as authentic."],
        ["Multipart FormData", "HTTP encoding used to upload media plus engine/source fields."],
        ["WebRTC / getUserMedia", "Browser standards used to request a camera MediaStream with user permission."],
        ["MediaRecorder", "Browser interface used to turn the live MediaStream into a three-second WebM Blob."],
    ], [2100, 7260], 9.2)

    doc.add_heading("Appendix A. File-Level Implementation Map", level=1)
    add_table(doc, ["File or directory", "Key implementation facts"], [
        ["index.html", "Dashboard structure, model selector, live and MP4 sources, viewport, verdict metrics, three sample buttons."],
        ["style.css", "Cyber-forensics visual language, responsive grids, scanner animation, terminal/result layers, red/secure states."],
        ["app.js", "State machine, camera permission, three-second recording, file handling, sample loading, POST request, calibration, verdict."],
        ["detector_server.py", "HTTP server, multipart parsing, temporary video handling, Meso4, FFmpeg sampling, remote routing, JSON."],
        [".models/meso4_best.pth", "Approximately 115 KB local Meso4 checkpoint."],
        ["NSU_Demo_Videos", "Three actual known-synthetic MP4s displayed as verification samples."],
        ["fetch_test_media.py", "Generates three stylized illustrative MP4s with Pillow/imageio. These are not the actual deepfake sample library."],
        ["README.md", "Local run command, detector explanation, Model 2 token setup, sample workflow."],
        [".deps/DeepfakeBench", "Reference implementation/configuration source used during model integration."],
        [".vendor", "Locally provided media dependencies, including imageio/FFmpeg support."],
    ], [2450, 6910], 9.2)
    doc.add_heading("Appendix B. Active Request Fields", level=1)
    add_table(doc, ["Field", "Values", "Meaning"], [
        ["media", "MP4 or WebM file", "Uploaded evidence or recorded live clip"],
        ["engine", "model1 / model2", "Selects local Meso4 or optional second route"],
        ["source_type", "live_video or omitted", "Marks camera recording for UI calibration and metadata"],
    ], [1700, 2500, 5160])
    doc.add_heading("Appendix C. Sources and Technical References", level=1)
    sources = [
        "Afchar, D., Nozick, V., Yamagishi, J., & Echizen, I. MesoNet: a Compact Facial Video Forgery Detection Network. https://arxiv.org/abs/1809.00888",
        "Yan, Z. et al. DeepfakeBench: A Comprehensive Benchmark of Deepfake Detection. https://arxiv.org/abs/2307.01426",
        "DeepfakeBench official repository. https://github.com/SCLBD/DeepfakeBench",
        "MDN Web Docs: MediaDevices.getUserMedia(). https://developer.mozilla.org/en-US/docs/Web/API/MediaDevices/getUserMedia",
        "MDN Web Docs: MediaRecorder. https://developer.mozilla.org/en-US/docs/Web/API/MediaRecorder",
        "Hive documentation: Image and Video Detection. https://docs.thehive.ai/docs/ai-image-and-video-detection",
        "Project source code and measured local test outputs in /Users/nusaibaalmollick/deepfake 2.0, inspected 13 July 2026.",
    ]
    for source in sources:
        add_bullet(doc, source, bullet_id)
    add_callout(doc, "Version note", "This document describes the repository state inspected on 13 July 2026. If the code changes, update the architecture, calibration, test results, and deployment sections before presentation.")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
